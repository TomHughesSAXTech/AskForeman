"""
Enhanced Document Converter for Azure Function
Supports: PDF, DOCX, TXT, CSV, XLSX, and more
"""

import logging
import io
import re
from typing import Optional

# PDF handling
import PyPDF2

# Word document handling
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not installed. DOCX support disabled.")

# Excel handling
try:
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False
    logging.warning("openpyxl not installed. Excel support disabled.")

# CSV handling
import csv

# Image OCR handling
try:
    from PIL import Image
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    logging.warning("PIL/pytesseract not installed. OCR support disabled.")


def extract_text_from_file(file_content: bytes, file_name: str, mime_type: str = None) -> str:
    """
    Extract text from various file types based on extension and MIME type.
    """
    # Determine file type from extension
    file_ext = file_name.lower().split('.')[-1] if '.' in file_name else ''
    
    # Override with MIME type if provided
    if mime_type:
        mime_lower = mime_type.lower()
    else:
        mime_lower = ''
    
    try:
        # PDF Files
        if file_ext == 'pdf' or 'pdf' in mime_lower:
            return extract_pdf_text(file_content)
        
        # Word Documents
        elif file_ext in ['docx', 'doc'] or 'wordprocessingml' in mime_lower or 'msword' in mime_lower:
            return extract_docx_text(file_content)
        
        # Excel Files
        elif file_ext in ['xlsx', 'xls'] or 'spreadsheet' in mime_lower or 'excel' in mime_lower:
            return extract_excel_text(file_content)
        
        # CSV Files
        elif file_ext == 'csv' or 'csv' in mime_lower:
            return extract_csv_text(file_content)
        
        # Plain Text Files
        elif file_ext in ['txt', 'md', 'log', 'json', 'xml', 'html', 'htm', 'rtf'] or 'text/' in mime_lower:
            return extract_plain_text(file_content)
        
        # Image Files (OCR)
        elif file_ext in ['png', 'jpg', 'jpeg', 'tiff', 'bmp'] or 'image/' in mime_lower:
            return extract_image_text(file_content)
        
        # Default: try as plain text
        else:
            logging.warning(f"Unknown file type: {file_ext} / {mime_type}. Attempting text extraction.")
            return extract_plain_text(file_content)
            
    except Exception as e:
        logging.error(f"Error extracting text from {file_name}: {str(e)}")
        # Final fallback
        return extract_plain_text(file_content)


def extract_pdf_text(pdf_content: bytes) -> str:
    """Extract text from PDF using PyPDF2."""
    try:
        pdf_file = io.BytesIO(pdf_content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        text = ""
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            page_text = page.extract_text()
            
            # Clean up text
            page_text = re.sub(r'\s+', ' ', page_text)
            page_text = page_text.strip()
            
            text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
        
        return text.strip()
    except Exception as e:
        logging.error(f"PDF extraction failed: {str(e)}")
        return ""


def extract_docx_text(docx_content: bytes) -> str:
    """Extract text from DOCX files."""
    if not DOCX_AVAILABLE:
        logging.warning("DOCX support not available. Install python-docx.")
        return extract_plain_text(docx_content)
    
    try:
        doc_file = io.BytesIO(docx_content)
        doc = Document(doc_file)
        
        text = []
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text.append(' | '.join(row_text))
        
        return '\n\n'.join(text)
    except Exception as e:
        logging.error(f"DOCX extraction failed: {str(e)}")
        return ""


def extract_excel_text(excel_content: bytes) -> str:
    """Extract text from Excel files."""
    if not EXCEL_AVAILABLE:
        logging.warning("Excel support not available. Install openpyxl.")
        return ""
    
    try:
        excel_file = io.BytesIO(excel_content)
        workbook = openpyxl.load_workbook(excel_file, read_only=True, data_only=True)
        
        text = []
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            text.append(f"\n--- Sheet: {sheet_name} ---\n")
            
            for row in sheet.iter_rows(values_only=True):
                # Filter out empty cells and convert to strings
                row_values = [str(cell) if cell is not None else '' for cell in row]
                if any(row_values):  # Skip completely empty rows
                    text.append(' | '.join(row_values))
        
        workbook.close()
        return '\n'.join(text)
    except Exception as e:
        logging.error(f"Excel extraction failed: {str(e)}")
        return ""


def extract_csv_text(csv_content: bytes) -> str:
    """Extract text from CSV files."""
    try:
        # Try to decode with different encodings
        for encoding in ['utf-8', 'latin-1', 'cp1252']:
            try:
                text_content = csv_content.decode(encoding)
                break
            except UnicodeDecodeError:
                continue
        else:
            text_content = csv_content.decode('utf-8', errors='ignore')
        
        csv_file = io.StringIO(text_content)
        reader = csv.reader(csv_file)
        
        text = []
        for row in reader:
            if row:  # Skip empty rows
                text.append(' | '.join(row))
        
        return '\n'.join(text)
    except Exception as e:
        logging.error(f"CSV extraction failed: {str(e)}")
        return extract_plain_text(csv_content)


def extract_plain_text(content: bytes) -> str:
    """Extract plain text with multiple encoding attempts."""
    # Try different encodings
    encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
    
    for encoding in encodings:
        try:
            text = content.decode(encoding)
            # Basic cleanup
            text = re.sub(r'\r\n', '\n', text)  # Normalize line endings
            text = re.sub(r'\n{3,}', '\n\n', text)  # Reduce multiple newlines
            return text.strip()
        except UnicodeDecodeError:
            continue
    
    # Final fallback with error handling
    return content.decode('utf-8', errors='ignore').strip()


def extract_image_text(image_content: bytes) -> str:
    """Extract text from images using OCR."""
    if not OCR_AVAILABLE:
        logging.warning("OCR support not available. Install pillow and pytesseract.")
        return ""
    
    try:
        image = Image.open(io.BytesIO(image_content))
        
        # Convert to RGB if necessary
        if image.mode != 'RGB':
            image = image.convert('RGB')
        
        # Extract text using OCR
        text = pytesseract.image_to_string(image)
        
        # Clean up
        text = re.sub(r'\s+', ' ', text)
        text = text.strip()
        
        return text
    except Exception as e:
        logging.error(f"OCR extraction failed: {str(e)}")
        return ""


# Integration with your existing Azure Function
def extract_document_text(file_content: bytes, file_name: str, mime_type: Optional[str] = None) -> str:
    """
    Main entry point for document text extraction.
    This replaces the extract_pdf_text function in your Azure Function.
    """
    logging.info(f"Extracting text from: {file_name} (MIME: {mime_type})")
    
    text = extract_text_from_file(file_content, file_name, mime_type)
    
    if not text:
        logging.warning(f"No text extracted from {file_name}")
        return f"[No readable text content found in {file_name}]"
    
    logging.info(f"Successfully extracted {len(text)} characters from {file_name}")
    return text
